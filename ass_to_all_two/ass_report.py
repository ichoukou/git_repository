# -*- coding: utf-8 -*-

import json
import ass_base
import ass_config
from reportlab.pdfgen import canvas
#用于定位的inch库，inch将作为我们的高度宽度的单位
from reportlab.lib.units import inch 
import reportlab.lib.fonts              
#canvas画图的类库
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib import fonts,colors 
from reportlab.pdfbase.ttfonts import TTFont 
# pdfmetrics.registerFont(ttfonts.TTFont("zhongwen", "simsun.ttc"))
from reportlab.pdfbase import pdfmetrics  
from reportlab.pdfbase.cidfonts import UnicodeCIDFont  
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))  
from reportlab.lib.pagesizes import letter, A4  
from reportlab.lib.styles import ParagraphStyle,PropertySet
from reportlab.lib.styles import getSampleStyleSheet
#from reportlab.platypus import Paragraph  
from reportlab.lib.enums import *  
from reportlab.lib.colors import *  
#from reportlab.lib.styles import getSampleStyleSheet  
from reportlab.platypus import Paragraph, SimpleDocTemplate, PageBreak  
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.rl_config import defaultPageSize
from docx.shared import RGBColor

from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
from reportlab.platypus import *
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY
from docx.shared import Inches
from itemManager import ItemManager
from itemManager import RiskCount
import subprocess
from threading import Timer
import os
import sys

from pyPdf import PdfFileWriter, PdfFileReader
from docx.shared import Pt

from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
reportlab.rl_config.warnOnMissingFontGlyphs = 0
#pdfmetrics.registerFont(TTFont('song', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc'))
#pdfmetrics.registerFont(TTFont('hei', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc')

from docx import Document

fontsName = 'msyh'
islinux = (sys.platform != "win32")
if islinux:
    fontsPath = ass_config.pinggu_dir + '/tool_linux/STZHONGS.TTF'
else:
    fontsPath = ass_config.pinggu_dir + '/tool_win/STZHONGS.TTF'


pdfmetrics.registerFont(TTFont(fontsName, fontsPath)) 
fonts.addMapping(fontsName, 0, 0, fontsName)
fonts.addMapping(fontsName, 0, 1, fontsName)

import sys, time, traceback

class AssEncoder(json.JSONEncoder):
    def default(self, obj):
        #convert object to a dict
        d = {}
        #d['__class__'] = obj.__class__.__name__
        #d['__module__'] = obj.__module__
        d.update(obj.__dict__)
        return d
 
class AssDecoder(json.JSONDecoder):
    def __init__(self):
        json.JSONDecoder.__init__(self, object_hook = self.dict2object)
    def dict2object(self,d):
        #convert dict to object
        if'__class__' in d:
            class_name = d.pop('__class__')
            module_name = d.pop('__module__')
            module = __import__(module_name)
            class_ = getattr(module, class_name)

            args = dict((key.encode('utf-8'), value) for key, value in d.items()) #get args
            inst = class_(**args) #create new instance
        else:
            inst = d
        return inst


class Basic:
    def __init__(self):
        self.appName = ''
        self.appVersion = ''
        self.packageName = ''

    # def __init__(self, obj):
    #     self.appName = obj.appName
    #     self.appVersion = obj.appVersion
    #     self.packageName = obj.packageName

class Result:
    def __init__(self):
        self.lastValue = 0
        self.evaAdvice = ''

    # def __init__(self, obj):
    #     self.lastValue = obj.lastValue
    #     self.evaAdvice = obj.evaAdvice
        
class List:
    def __init__(self):
        self.items = []
        self.type = ''

    # def __init__(self, obj):
    #     self.items = obj.items
    #     self.type = obj.type

class Items:
    def __init__(self):
        self.checkMethod = ''
        self.checkName = ''
        self.checkResult = ''
        self.relatedCode = ''
        self.resultComment = ''

    # def __init__(self, obj):
    #     self.checkMethod = obj.checkMethod
    #     self.checkName = obj.checkName
    #     self.checkResult = obj.checkResult
    #     self.relatedCode = obj.relatedCode
    #     self.resultComment = obj.resultComment

class Report:
    def __init__(self):
        self.basic = Basic()
        self.list = [List()]
        self.result = Result()

pdfmetrics.registerFont(TTFont('simfang', ass_config.pinggu_dir + '/assets/simfang.ttf'))
fonts.addMapping('simfang', 0, 0, 'simfang')
fonts.addMapping('simfang', 0, 1, 'simfang')
fonts.addMapping('simfang', 1, 0, 'simfang')
fonts.addMapping('simfang', 1, 1, 'simfang')

PAGE_HEIGHT = defaultPageSize[1]
PAGE_WIDTH = defaultPageSize[0]

def drawFoot(canvas, doc):
    canvas.saveState()
    canvas.setFont('simfang', 9)
    canvas.drawCentredString(PAGE_WIDTH / 4.0 * 3.0, 30.0, u'第 %d 页' % (doc.page))
    p = canvas.beginPath()
    p.moveTo(50, 40)
    p.lineTo(PAGE_WIDTH - 50, 40)
    p.close()
    canvas.drawPath(p)
    canvas.restoreState()

def drawHeader(canvas, doc):
    canvas.saveState()
    LOGO = ass_config.pinggu_dir + '/assets/logobe.png'  #147*46
    hSize, vSize = 105, 25
    offset = 30
    canvas.drawImage(LOGO, 70, PAGE_HEIGHT-vSize-6.5-offset, hSize, vSize)
    canvas.setFont('simfang', 21)
    canvas.drawString(220, PAGE_HEIGHT-21.0-offset, u'移动APP安全检测报告')
    canvas.setFont('simfang', 14)
    canvas.drawString(220, PAGE_HEIGHT-21.0-14.0-offset, u'鼎源科技APP安全检测扫描平台（www.appbesafe.com）')
    p = canvas.beginPath()
    p.moveTo(50, PAGE_HEIGHT-40.0-offset)
    p.lineTo(PAGE_WIDTH-50, PAGE_HEIGHT-40.0-offset)
    p.close()
    canvas.drawPath(p)
    canvas.restoreState()

def onTitlePage(canvas, doc):
    drawHeader(canvas, doc)

def onPage(canvas, doc):
    drawHeader(canvas, doc)
    drawFoot(canvas, doc)


class MyDocTemplate(BaseDocTemplate):
    def __init__(self, filename, **kw):
        self.allowSplitting = 0
        BaseDocTemplate.__init__(self, filename, **kw)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id='normal')
        self.addPageTemplates([
                              #PageTemplate(id='Title_Page', frames=[frame], onPage=onTitlePage),
                              #PageTemplate(id='Table_of_Contents', frames=[frame], onPage=onPage),
                              PageTemplate(id='Body', frames=[frame], onPage=onPage)])
#检测报告生成模块
class AssReport(ass_base.AssI18n):
    def __init__(self):
        self.modules = []
        self.apk_file = ''
        self.language = ''
        self.apk_info = ''
        self.inited = False
        self.title = 'APP风险评估报告'  #标题
        self.manager = ItemManager()
        self.manager.initData()
        #self.count = RiskCount(0, 0, 0)
        #self.count.getTotalCount(0)

    #设置检测项的结果和详情
    def setItem(self, itemID, detail, num=0):
        detail.strip()
        
        import chardet
        charInfo = chardet.detect(detail)
        detail = detail.decode(charInfo['encoding'])
        if len(detail) > 500:
            detail = detail[:500]

        if itemID[0] == '0':                #应用安全
            self.manager.updateAppItem(itemID, True, detail, num)
        elif itemID[0] == '1':              #源码安全
            self.manager.updateSourceItem(itemID, True, detail, num)
        elif itemID[0] == '2':              #数据安全
            self.manager.updateDataItem(itemID, True, detail, num)
        elif itemID[0] == '3':              #漏洞
            self.manager.updateVulItem(itemID, True, detail, num)
        elif itemID[0] == '4':              #恶意行为
            self.manager.updateHarmItem(itemID, True, detail, num)

    def setPermissionItem(self, itemID, permissionName):
        self.manager.updatePermissionItem(itemID, permissionName)

    def setBaseInfo(self, value, type):
        self.manager.updateBaseInfo(value, type)

    def getResult(self):
        return self.manager.getResult()

    #格式化检测信息，添加报告输出列表
    #[code 相关信息] [comment 问题描述] [result 等级] [name 检测说明] [method 检测方法]       
    def addItem(self, code, comment='', name='攻击检测', result='高危', method='全自动化分析'):
        items = Items()
        items.checkMethod = self.i18n(method)
        items.checkName = self.i18n(name)
        items.checkResult = self.i18n(result)
        MaxLen = 256
        if len(code) > MaxLen:
            items.relatedCode = code[0:MaxLen]
        else:
            items.relatedCode = code
        items.resultComment = self.i18n(comment)
        self.report.list[0].items.append(items)
        self.addResultValue(self.i18n(result))

    #格式化信息，添加到输出报告列表
    #[arr 目标数组] [comment 信息]   
    def addArrItem(self, arr, comment, cat='攻击检测'):
        if len(arr) > 0:
            self.addItem(','.join(arr), comment, cat)

    #添加检测安全度系数
    #[result 检测项风险等级]   
    def addResultValue(self, result):
        if result == self.i18n('高危'):
            self.report.result.lastValue += 1
        elif result == self.i18n('危险'):
            self.report.result.lastValue += 0.5

    #添加检测安全度系数
    #[result 检测项风险等级]
    def updateResult(self):
        self.result = self.manager.getResult()

    def setValidFlag(self, v):
        self.manager.setValidFlag(v)

    def getValidFlag(self):
        return self.manager.getValidFlag()

    def setSignFlag(self, v):
        self.manager.setSignFlag(v)

    def getSignFlag(self):
        return self.manager.getSignFlag()

    def getReport(self):
        self.updateResult()
        if not hasattr(self, "report"):
            return ''
        return AssEncoder(indent = 2, ensure_ascii = False).encode(self.report)

    #将长字符转换为列表
    #[str_c 字符] [pos 字符长度限制]
    def strToList(self, str_c, pos=32):
        str = str_c.decode('utf-8')
        #s.decode('utf-8')[x*y:x*(y+1)].encode('utf-8')
        lstr = []
        str = str.replace(unicode("\n"), unicode(""))
        str = str.replace(unicode("\t"), unicode(""))
        str = str.replace(unicode("\r"), unicode(""))
        strlen = len(str)
        while strlen > 0:
            if strlen > pos:
                lstr.append(str[:pos].encode('utf-8'))
                str = str[pos:]
            else:
                lstr.append(str.encode('utf-8'))
                str = ''
            strlen = len(str)
        return lstr

    #创建pdf文件输出
    #[output pdf输出文件] [title 标题]
    def create_pdf(self, output, title):
        #表格输出
        #self.pdf_body_table(output, title)
        self.pdf_body_out(output)

        #添加文件水印 
        self.add_watermark(output)
        #普通输出
        #c = canvas.Canvas(output) 
        #self.pdf_head(c, title)
        #self.pdf_body(c,title)

    def pdf_body_out(self, output):
        story = []
        doc = MyDocTemplate(output.encode('utf-8'))
        self.manager.genStandardsTable(story)
        self.manager.genAppInfoTable(story)
        self.manager.genAbsTable(story)
        self.manager.genChartTable(story, output + '@')
        self.manager.genSummaryTable(story)
        self.manager.genPermissionInfo(story)
        self.manager.genBody(story)
        self.manager.genLevelDescTable(story)
        doc.build(story)

    #创建pdf文件内容表格输出
    #[output pdf输出文件] [title 标题] 
    def pdf_body_table(self, output, title):
        stylesheet = getSampleStyleSheet()
        elements = []
        doc = SimpleDocTemplate(output.encode('utf-8'))

        stylesheet.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
        stylesheet['Justify'].fontName = fontsName
        elements.append(flowables.Preformatted('', stylesheet['Justify']))
        elements.append(Spacer(1,12))
       
        ts = [('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black), ('BOX', (0, 0), (-1, -1), 0.25, colors.black), ('FONT', (0, 0), (-1, -1), fontsName)]

        #tempstr="安全评估结论" #打印安全评估结论
        #tempstr = "评估结果:"+self.report.result.evaAdvice
        #tempstr = "危险系数:"+str(self.report.result.lastValue)
        #tempstr ="应用基本信息"
        #tempstr ="应用名称:"+self.report.basic.appName
        #tempstr = "应用版本:"+str(self.report.basic.appVersion)
        #tempstr = "应用包名:"+self.report.basic.packageName
        elements.append(self.pdf_table_null(ts, '安全评估结论'))#安全评估结论
        elements.append(self.pdf_table_format('评估结果:', self.report.result.evaAdvice,ts))
        elements.append(self.pdf_table_format('危险系数:', str(self.report.result.lastValue), ts))#危险系数
        
        elements.append(self.pdf_table_null(ts, '应用基本信息'))#应用基本信息
        elements.append(self.pdf_table_format('应用名称:', self.report.basic.appName, ts))
        elements.append(self.pdf_table_format('应用版本:', self.report.basic.appVersion, ts))
        elements.append(self.pdf_table_format('应用包名:', self.report.basic.packageName, ts))
        elements.append(self.pdf_table_null(ts, '详细检测项'))#空表间隔

        count_check = 0;#统计风险总数
        for i in range(len(self.report.list)):
            elements.append(self.pdf_table_format('检测项:', self.report.list[i].type, ts))
            elements.append(self.pdf_table_null(ts))#空表间隔
            count_check = count_check + len(self.report.list[i].items)  #统计
            for j in range(len(self.report.list[i].items)):
                data = []
                elements.append(self.pdf_table_format('评估方法:', self.report.list[i].items[j].checkMethod, ts))
                elements.append(self.pdf_table_format('评估项:', self.report.list[i].items[j].checkName, ts))
                elements.append(self.pdf_table_format('评估建议:', self.report.list[i].items[j].checkResult, ts))
                elements.append(self.pdf_table_format('详细描述:', self.report.list[i].items[j].resultComment, ts))
                elements.append(self.pdf_table_format('相关代码:', self.report.list[i].items[j].relatedCode, ts))
                elements.append(self.pdf_table_null(ts))#空表间隔

        elements.append(self.pdf_table_format('总共发现风险:',str(count_check) + "项",ts))#发现风险数量
        doc.build(elements)

    #"""给指定PDF文件文件加上水印 
    #pdf_file - 要加水印的源PDF文件 
    #"""  
    def add_watermark(self,pdf_file):  
        #获取pdf样式样本
        pdf_template = PdfFileReader(file(ass_config.pdf_template, 'rb')) 

        pdf_output = PdfFileWriter()  
        #一、添加首页
        pdf_output.addPage(pdf_template.getPage(0))
        #二、添加水印
        input_stream = file(pdf_file, 'rb')  
        pdf_input = PdfFileReader(input_stream)  
      
        # PDF文件被加密了  
        if pdf_input.getIsEncrypted():  
            print '该PDF文件被加密了.'  
            # 尝试用空密码解密  
            try:  
                pdf_input.decrypt('')  
            except Exception, e:  
                print '尝试用空密码解密失败.'  
                return False  
            else:  
                print '用空密码解密成功.'  
  
        # 获取PDF文件的页数   
        pageNum = pdf_input.getNumPages()   
        # 给每一页打水印   
        for i in range(pageNum):  
            page = pdf_input.getPage(i)  
            # 加水印  
            # try:
            #     page.mergePage(pdf_template.getPage(1))
            # except:
            #     print traceback.print_exc()
            #     self.ERROR("watermark add error page %d" % i)
            pdf_output.addPage(page)  
  
        # 最后输出文件   
        output_stream = file(os.path.join(os.path.dirname(pdf_file), os.path.basename(pdf_file)[0:-4] + '-final.pdf'), 'wb')
        pdf_output.write(output_stream)   
        output_stream.close()  
        input_stream.close() 
        #删除原文件
        os.remove(pdf_file)

    #定义表格格式
    #[内容 data1] [内容 data2]
    def pdf_table_format(self, data1, data2, ts):
        dataList = []
        strList = self.strToList(data2, 40)
        dataList.append([data1, self.pdf_table_list2str(strList)])
        height = len(strList) * 0.2
        if height == 0:
            height = 0.2
        wlist = [inch, 6*inch]
        table = Table(dataList, wlist, height*inch, ts)
        return table

    #定义空表格格式
    def pdf_table_null(self,ts,data = ''):
        dataList = []
        dataList.append([data])
        table = Table(dataList, 7*inch, 0.4*inch, ts)
        return table

    #将列表转换为字符
    #[strList 源列表]
    def pdf_table_list2str(self,strList):
        str =''
        for j in range(len(strList)):
           str = str + '\n' + strList[j]
        return str

    #设置字体
    #[size 字体大小]
    def pdf_fontSet(self,c,size):
        c.setFont(fontsName, size) 

    #绘制文本
    #[x 横坐标] [y 横坐标] [text 文本内容]
    def pdf_drawText(self,c,x,y,text):
        try:
           c.drawString(x, y, text)
        except:
           xx = x/inch
           yy = y/inch
           te = text
           print("[except,%d,%d,%s]"%(xx,yy, te))
           print trace_back()

    #绘制文本
    #[x 横坐标] [y 横坐标] [text 文本内容]
    def pdf_drawStr(self,c,textList,restInch):
        #假设每行占据0.2inch,计算此段需要占据多少
        size = 0.2
        try:
            count = len(textList)
            restCount = restInch / size
            restCount = int(restCount)
            if count > restCount:
                if restCount >= 1:
                    for i in range(restCount):
                        self.pdf_drawText(c, 1*inch, restInch*inch, textList[i])
                        restInch=restInch-size
                c.showPage()#重新分页
                self.pdf_head(c, self.title)
                restInch = 10
                restInch = self.pdf_drawStr(c, textList[restCount:], restInch)
            else:
                for j in range(count):
                    self.pdf_drawText(c, 1*inch, restInch*inch, textList[j])
                    restInch = restInch-size
            return restInch
        except:
            return restInch

    #pdf标题
    #[headtext 标题]
    def pdf_head(self,c,headtext):
        self.pdf_fontSet(c,9)
        c.drawCentredString(4*inch,10.5*inch,headtext)
        #c.drawString(4*inch, 10.5*inch, headtext)
        c.rect(1*inch, 10.3*inch, 6.5*inch, 0.12*inch,fill=1) 

    #pdf内容
    #[title 标题]
    def pdf_body(self,c,title): 
        self.pdf_fontSet(c,13);   
        tempstr="安全评估结论" #打印安全评估结论
        c.drawString(1*inch, 9.8*inch, tempstr)
        self.pdf_fontSet(c,9);#评估结果
        tempstr = "评估结果:" + self.report.result.evaAdvice
        c.drawString(1*inch, 9.6*inch, tempstr)
        tempstr = "危险系数:"+str(self.report.result.lastValue)
        c.drawString(1*inch, 9.4*inch, tempstr)#危险系数
        c.line(1*inch, 9.2*inch, 7.5*inch, 9.2*inch) 
        self.pdf_fontSet(c,13);
        tempstr ="应用基本信息"
        c.drawString(1*inch, 9*inch, tempstr)#应用基本信息
        self.pdf_fontSet(c,9);
        tempstr ="应用名称:"+self.report.basic.appName
        c.drawString(1*inch, 8.8*inch, tempstr)#应用名称
        tempstr = "应用版本:"+str(self.report.basic.appVersion)
        c.drawString(1*inch, 8.6*inch, tempstr)#应用版本
        tempstr = "应用包名:"+self.report.basic.packageName
        c.drawString(1*inch, 8.4*inch, tempstr)#应用包名
        c.line(1*inch, 8.2*inch, 7.5*inch, 8.2*inch)
        self.pdf_fontSet(c,13);
        tempstr ="详细检测项" 
        c.drawString(1*inch, 8*inch, tempstr)#详细检测项
        self.pdf_fontSet(c,9); 
        tempf = 8
        key = 4
        relateCodes = []
        for i in range(len(self.report.list)):
            tempf=8-0.2
            tempstr ="检测项:"+self.report.list[i].type
            tempf=self.pdf_drawStr(c,self.strToList(tempstr),tempf)#详细检测项
            tempf-=0.2
            for j in range(len(self.report.list[i].items)):
                tempstr ="评估方法:"+self.report.list[i].items[j].checkMethod
                tempf=self.pdf_drawStr(c,self.strToList(tempstr),tempf)#评估方法
                tempstr ="评估项:"+self.report.list[i].items[j].checkName
                tempf=self.pdf_drawStr(c,self.strToList(tempstr),tempf)#评估项
                tempstr ="评估建议:"+self.report.list[i].items[j].checkResult
                tempf=self.pdf_drawStr(c,self.strToList(tempstr),tempf)#评估建议
                tempstr ="详细描述:"+self.report.list[i].items[j].resultComment
                tempf=self.pdf_drawStr(c,self.strToList(tempstr),tempf)#详细描述
                tempstr ="相关代码:"+self.report.list[i].items[j].relatedCode
                tempf=self.pdf_drawStr(c,self.strToList(tempstr,100),tempf)#相关代码
                c.line(1*inch, tempf*inch, 7.5*inch, tempf*inch)
                tempf=tempf-0.2
        c.save() 
        
    #保存pdf报告
    #[apk_file 保存apk文件]                             
    def saveReport(self, apk_file, module=''):
        #out_file = apk_file+'-report'+".json";
        #out_txt_file = apk_file+'-report'+'.txt';
        #out_pdf_file = apk_file + '-report-' + self.manager.appName + ".pdf"
        out_pdf_file = apk_file + '-report' + ".pdf"
        out_json_file = apk_file + '-report' + ".json"
        out_word_file = apk_file + '-report' + ".docx"
        self.manager.toJsonFile(out_json_file)
        #保存json
        #ass_base.write_file(out_file, self.getReport())

        # #保存txt
        # tempstr="安全评估结论\n"
        # tempstr += "评估结果："+self.report.result.evaAdvice+"\n"
        # tempstr += "危险系数："+str(self.report.result.lastValue)+"\n\n"
        # tempstr +="应用基本信息\n"
        # tempstr +="应用名称："+self.report.basic.appName+"\n"
        # tempstr += "应用版本："+str(self.report.basic.appVersion)+"\n"
        # tempstr += "应用包名："+self.report.basic.packageName+"\n\n"
        # tempstr +="详细检测项\n"
        # for i in range(len(self.report.list)):
        #     tempstr +="检测项："+self.report.list[i].type+"\n\n"
        #     for j in range(len(self.report.list[i].items)):
        #         tempstr +="评估方法:"+self.report.list[i].items[j].checkMethod+"\n"
        #         tempstr +="评估项："+self.report.list[i].items[j].checkName+"\n"
        #         tempstr +="评估建议："+self.report.list[i].items[j].checkResult+"\n"
        #         tempstr +="详细描述："+self.report.list[i].items[j].resultComment+"\n"
        #         tempstr +="相关代码：\n"+self.report.list[i].items[j].relatedCode+"\n\n"
        # ass_base.write_file(out_txt_file, tempstr)
        #保存pdf
        #self.create_pdf(out_pdf_file, self.title)
        self.result = self.getResult()
        self.create_word(out_word_file)
        self.do_cmd(ass_config.cmd_convertReport + " " + out_word_file +  " " + out_pdf_file)

        #self.readReport(out_file)

    def to_kill(p):
       if not p==None:
        p.terminate()

    def do_cmd(self, cmdline, print_out=True, timeout=0):#执行cmd
       #cmdline = unicode(cmdline,"utf8")
       print("++++ [ CMD ] %s  " % cmdline)
       p = subprocess.Popen(cmdline, bufsize = 0, shell = True, universal_newlines = True, stdin = subprocess.PIPE, stdout = subprocess.PIPE, stderr = subprocess.STDOUT)
       if timeout > 0:
           t = Timer(timeout, to_kill(p))
           t.start()
       out = ''
       err = ''
       try:
           out,err = p.communicate()
       except:
           out = ''
           err = ''

       if print_out:
           #全部输出
           print(out)
       else:
           #截取输出
           if len(out)> 200:
               print(out[0:200])
           else:
               print(out)
       if timeout > 0:
           t.cancel()
       return out

    def genAutoResultList(self, items, resultList):
        if items != None and len(items) > 0:
            ids = self.manager.sortBySortId(items, 0)
            for key in ids:
                pItem = items[key]
                if pItem.type == 0:
                    resultList.append(pItem.result)

    def genAutoCount(self, items):#counttt
        lowCount, midCount, highCount = 0, 0, 0
        if items != None and len(items) > 0:
            ids = self.manager.sortBySortId(items, 0)
            for key in ids:
                pItem = items[key]
                if pItem.type == 0 and pItem.state:
                    print pItem.name + '**********'+pItem.level + '*********'
                    if pItem.score == 1:
                        lowCount += pItem.account
                    elif pItem.score == 2:
                        midCount += pItem.account
                    elif pItem.score == 3:
                        highCount += pItem.account
        return [lowCount, midCount, highCount]

    def genResultListWithDetail(self, items, resultList, type):
        if items != None and len(items) > 0:
            ids = self.manager.sortBySortId(items, 0)
            for key in ids:
                pItem = items[key]
                if pItem.type == type:
                    resultList.append((pItem.result, pItem.detail, pItem.name))

    def setCellFont(self, cell, text, style, ali = 0):
        p = cell.paragraphs[0]
        p.add_run(text, style = style).bold = True
        if ali  ==  0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #paragraph = cell.paragraphs[0]


    def genRow(self, sRowCount, table, items, clsName):
        rowCount = sRowCount

        if items != None and len(items) > 0:
            ids = self.manager.sortBySortId(items, 0)
            for key in ids:
                pItem = items[key]
                if pItem.type == 0 and pItem.result == u'存在风险':
                    r = table.add_row()
                    if rowCount == sRowCount:
                        r.cells[0].text = clsName
                    r.cells[1].text = pItem.name
                    r.cells[2].text = pItem.level
                    r.cells[3].text = pItem.result
                    print r.cells[1].text + "========" + str(rowCount)
                    rowCount += 1
        if rowCount > sRowCount:
            cellBegin = table.cell(sRowCount, 0)
            cellBegin.merge(table.cell(rowCount - 1, 0))
        return rowCount

    def create_word_short(self,out_word_file):
        doc = Document('assets/tempshort.docx')

        tab_1 = doc.tables[4]
        rowCount = self.genRow(2, tab_1, self.manager.appItems, u'应用安全')
        rowCount = self.genRow(rowCount, tab_1, self.manager.sourceItems, u'源码安全')
        rowCount = self.genRow(rowCount, tab_1, self.manager.dataItems, u'数据安全')
        rowCount = self.genRow(rowCount, tab_1, self.manager.vulItems, u'漏洞安全')
        rowCount = self.genRow(rowCount, tab_1, self.manager.harmItems, u'恶意行为')
        doc.save(out_word_file)

    def set_normal_info(self, doc, num):
        doc = doc
        i = num
        tab_1 = doc.tables[i]
        self.setCellFont(tab_1.cell(1, 1), self.manager.appName, "item", 1)
        self.setCellFont(tab_1.cell(2, 1), self.manager.appPackageName, "item", 1)
        self.setCellFont(tab_1.cell(3, 1), self.manager.appVersion, "item", 1)
        self.setCellFont(tab_1.cell(4, 1), self.manager.fileSize, "item", 1)
        self.setCellFont(tab_1.cell(5, 1), self.manager.md5.strip(), "item", 1)
        self.setCellFont(tab_1.cell(6, 1), self.manager.checkEndTime, "item", 1)
        self.setCellFont(tab_1.cell(7, 1), self.manager.checkShell, "item", 1)

    def list_info(self, doc, num):
        resultList = []
        self.genAutoResultList(self.manager.appItems, resultList)
        self.genAutoResultList(self.manager.sourceItems, resultList)
        self.genAutoResultList(self.manager.dataItems, resultList)
        self.genAutoResultList(self.manager.vulItems, resultList)
        self.genAutoResultList(self.manager.harmItems, resultList)

        tab_1 = doc.tables[num]
        for i in range(len(resultList)):
            cell = tab_1.cell(2 + i, 3)
            if resultList[i] == u'存在风险':
                self.setCellFont(cell, u'存在风险', "unsafe")
            else:
                self.setCellFont(cell, u'不存在风险', "safe")


    def list_detail(self, doc, num):
        items = self.manager.permissionItems
        ids = None
        if items != None and len(items) > 0:
            ids = self.manager.sortBySortId(items, 1)
            tab_1 = doc.tables[num]
            print len(tab_1.columns)
            print len(tab_1.rows)

            for row in range(len(tab_1.rows)):
                if row >= 2:
                    self.setCellFont(tab_1.cell(row, 3), u"不存在风险", "safe", 1)
                    for key in ids:
                        pItem = items[key]
                        if pItem.name in tab_1.cell(row, 1).text:
                            tab_1.cell(row, 3).paragraphs[0].clear()
                            #tab_1.cell(row, 3).add_paragraph("")
                            #tab_1.cell(row, 3).text = ''
                            self.setCellFont(tab_1.cell(row, 3), u"存在风险", "unsafe", 1)
                            break
        resultList = []
        self.genResultListWithDetail(self.manager.appItems, resultList, 0)
        self.genResultListWithDetail(self.manager.sourceItems, resultList, 0)
        self.genResultListWithDetail(self.manager.dataItems, resultList, 0)
        self.genResultListWithDetail(self.manager.vulItems, resultList, 0)
        self.genResultListWithDetail(self.manager.harmItems, resultList, 0)

        count = 0
        for i in resultList:
            tab_1 = doc.tables[count + 12]
            count += 1
            #print tab_1.cell(1, 0).text + "====>" + i[2]
            if i[0] == u'存在风险':
                self.setCellFont(tab_1.cell(5, 0), u'存在风险', "unsafe", 1)
            else:
                self.setCellFont(tab_1.cell(5, 0), u'不存在风险', "safe", 1)
            self.setCellFont(tab_1.cell(7, 0), i[1], "item", 1)

    def getAPPSummary(self, doc):
        flag = [0, 0, 0, 0]
        #应用安全的总结
        tab_1 = doc.tables[4]
        if self.manager.appItems != None and len(self.manager.appItems) > 0:
            sSummary = ""
            ids = self.manager.sortBySortId(self.manager.appItems, 0)
            for key in ids:
                pItem = self.manager.appItems[key]
                if pItem.type == 0 and pItem.name == u'WebView安全' and pItem.result == u'存在风险':
                    sSummary += u"谨慎使用WebView接口，防止WebView漏洞被恶意利用的风险；"
                if pItem.type == 0 and pItem.name == u'本地拒绝服务' and pItem.result == u'存在风险':
                    sSummary += u"在使用getAction时对返回值进行null检查后再使用；"
                if pItem.type == 0 and pItem.name == u'Service安全' and pItem.result == u'存在风险':
                    sSummary += u"只有应用自身使用的service应设置为私有；"
                if pItem.type == 0 and pItem.name == u'日志信息检查' and pItem.result == u'存在风险':
                    sSummary += u"关闭调试日志调用，或者确保日志的输出使用了正确的级别；"
                if pItem.type == 0 and pItem.name == u'Broadcast Receiver安全' and pItem.result == u'存在风险':
                    sSummary += u"加固APP防止Broadcast Receiver劫持和监听功能；"
                if pItem.type == 0 and pItem.name == u'Activity安全' and pItem.result == u'存在风险':
                    sSummary += u"验证目标Activity是否恶意app，以免受到intent欺骗，可用hash签名验证；"
                if pItem.type == 0 and pItem.name == u'动态注册广播' and pItem.result == u'存在风险':
                    sSummary += u"对接收来的广播进行验证，返回结果时需注意接收APP是否会泄露信息；"

            if len(sSummary) > 0:
                sSummary = u"   建议" + sSummary
                self.setCellFont(tab_1.cell(0, 0), sSummary, "item", 1)
                flag[0] = 1
            else:
                self.setCellFont(tab_1.cell(0, 0), u"   无", "item", 1)

        #源码安全的总结
        tab_1 = doc.tables[5]
        if self.manager.sourceItems != None and len(self.manager.sourceItems) > 0:
            sSummary = ""
            ids = self.manager.sortBySortId(self.manager.sourceItems, 0)
            for key in ids:
                pItem = self.manager.sourceItems[key]
                if pItem.type == 0 and pItem.name == u'反编译防范检测' and pItem.result == u'存在风险':
                    sSummary += u"采取防反编译功能，防止应用被反编译；"
                elif pItem.type == 0 and pItem.name == u'资源文件保护检查' and pItem.result == u'存在风险':
                    sSummary += u"使用具有资源文件防反编译功能的加密加固工具对APP中的资源文件进行加密；"
                elif pItem.type == 0 and pItem.name == u'主配置文件保护检查' and pItem.result == u'存在风险':
                    sSummary += u"使用具有AndroidManifest.xml文件防反编译功能的加密加固工具对APP中的AndroidManifest.xml文件进行加密；"
                elif pItem.type == 0 and pItem.name == u'完整性校验检查' and pItem.result == u'存在风险':
                    sSummary += u"采取必要的防反编译保护手段和二次打包校验措施；"
                elif pItem.type == 0 and pItem.name == u'硬编码-密码敏感词编码检查' and pItem.result == u'存在风险':
                    sSummary += u"利用无意义的变量名替换敏感词，做好对比记录；"
                elif pItem.type == 0 and pItem.name == u'安全编码规范-全反射调用检查' and pItem.result == u'存在风险':
                    sSummary += u"如果无需与JS交互，请删除对addJavascriptInterface函数的调用；"

            if len(sSummary) > 0:
                sSummary = u"   建议" + sSummary
                self.setCellFont(tab_1.cell(0, 0), sSummary, "item", 1)
                flag[1] = 1
            else:
                self.setCellFont(tab_1.cell(0, 0), u"   无", "item", 1)

        #数据安全的总结
        tab_1 = doc.tables[6]
        if self.manager.dataItems != None and len(self.manager.dataItems) > 0:
            sSummary = ""
            ids = self.manager.sortBySortId(self.manager.dataItems, 0)
            for key in ids:
                pItem = self.manager.dataItems[key]
                if pItem.type == 0 and pItem.name == u'存储数据检查' and pItem.result == u'存在风险':
                    sSummary += u"对本地文件的存储进行加密处理；"
                elif pItem.type == 0 and pItem.name == u'远程数据通讯协议' and pItem.result == u'存在风险':
                    sSummary += u"通信时使用加密信道；"

            if len(sSummary) > 0:
                sSummary = u"   建议" + sSummary
                self.setCellFont(tab_1.cell(0, 0), sSummary, "item", 1)
                flag[2] = 1
            else:
                self.setCellFont(tab_1.cell(0, 0), u"   无", "item", 1)

        #漏洞安全
        tab_1 = doc.tables[7]
        if self.manager.vulItems != None and len(self.manager.vulItems) > 0:
            sSummary = ""
            ids = self.manager.sortBySortId(self.manager.vulItems, 0)
            for key in ids:
                pItem = self.manager.vulItems[key]
                if pItem.type == 0 and pItem.name == u'zip文件目录遍历漏洞' and pItem.result == u'存在风险':
                    sSummary += u"使用zipEntry.getName()生成文件时，过滤上级目录字符串（../）；"
                if pItem.type == 0 and pItem.name == u'WebView组件忽略SSL证书验证错误' and pItem.result == u'存在风险':
                    sSummary += u"避免重写onReceivedSslError方法， 或者对于SSL证书错误问题按照业务场景判断，避免造成数据明文传输情况；"

            if len(sSummary) > 0:
                sSummary = u"   建议" + sSummary
                self.setCellFont(tab_1.cell(0, 0), sSummary, "item", 1)
                flag[3] = 1
            else:
                self.setCellFont(tab_1.cell(0, 0), u"   无", "item", 1)

        tab_1 = doc.tables[3]
        sSummary = u"   该APP的总体风险等级为 " + self.manager.getComment() + u"，总体建议"
        if flag[0] == 1:
            sSummary += u"加强对APP应用的基本攻击面的保护，注意相关组件和接口的安全设置，"
        if flag[1] == 1:
            sSummary += u"为APP添加保护壳，保证APK文件的完整性和安全性，防止算法和数据泄露，"
        if flag[2] == 1:
            sSummary += u"使用安全通信协议，保证其与服务器端的通信安全，确保本地数据经过加密措施保护"
        if flag[3] == 1:
            sSummary += u"对于常见的APP漏洞要采取补救措施或替换实现方案，避免漏洞残留，"

        sSummary += u"针对相应检测项进行对应的保护处理，让您的APP更加安全健壮。"
        self.setCellFont(tab_1.cell(0, 0), sSummary, "item", 1)


    def create_word(self, out_word_file):
        doc = Document('assets/temp.docx')

        print len(doc.tables)
        print doc.tables
        #设置基本信息
        self.set_normal_info(doc, 0)
        tab_1 = doc.tables[1]
        count1 = [0, 0, 0]
        count1 = map(lambda (a,b):a+b, zip(count1, self.manager.getRealCount(self.manager.appItems)))
        count1 = map(lambda (a,b):a+b, zip(count1, self.manager.getRealCount(self.manager.dataItems)))
        count1 = map(lambda (a,b):a+b, zip(count1, self.manager.getRealCount(self.manager.vulItems)))
        count1 = map(lambda (a,b):a+b, zip(count1, self.manager.getRealCount(self.manager.sourceItems)))
        count1 = map(lambda (a,b):a+b, zip(count1, self.manager.getRealCount(self.manager.harmItems)))

        self.setCellFont(tab_1.cell(2, 4), str(count1[2]), "item", 0)
        self.setCellFont(tab_1.cell(2, 3), str(count1[1]), "item", 0)
        self.setCellFont(tab_1.cell(2, 2), str(count1[0]), "item", 0)
        self.setCellFont(tab_1.cell(2, 1), str(count1[2] + count1[1] + count1[0]), "item", 0)
        self.setCellFont(tab_1.cell(4, 1), self.manager.getComment(), "unsafe", 0)
        print str(count1[0])+">>>>>>>>>>>>"+str(count1[1])+">>>>>>>>>>>>"+str(count1[2])

        count = [0, 0, 0]
        count = map(lambda (a,b):a+b, zip(count, self.genAutoCount(self.manager.appItems)))
        count = map(lambda (a,b):a+b, zip(count, self.genAutoCount(self.manager.dataItems)))
        count = map(lambda (a,b):a+b, zip(count, self.genAutoCount(self.manager.vulItems)))
        count = map(lambda (a,b):a+b, zip(count, self.genAutoCount(self.manager.sourceItems)))
        count = map(lambda (a,b):a+b, zip(count, self.genAutoCount(self.manager.harmItems)))
        #print str(count[0]) + ">>>>>>>>>" + str(count[1]) + ">>>>>>>" + str(count[2])
        self.setCellFont(tab_1.cell(3, 2), str(count[2]), "item", 0)
        self.setCellFont(tab_1.cell(3, 3), str(count[1]), "item", 0)
        self.setCellFont(tab_1.cell(3, 4), str(count[0]), "item", 0)
        self.setCellFont(tab_1.cell(3, 1), str(count[2] + count[1] + count[0]), "item", 0)

        tab_1 = doc.tables[2]
        tab_1.cell(1, 0).add_paragraph().add_run().add_picture(self.manager.genChartImg(out_word_file + '.png', 0), Inches(5))


        self.list_info(doc, 9)
        self.list_detail(doc, 11)
        self.getAPPSummary(doc)
        #self.genResultListWithDetail(self.manager.appItems, resultList, 1)
        #self.genResultListWithDetail(self.manager.sourceItems, resultList, 1)
        #self.genResultListWithDetail(self.manager.dataItems, resultList, 1)
        #self.genResultListWithDetail(self.manager.vulItems, resultList, 1)
        #self.genResultListWithDetail(self.manager.harmItems, resultList, 1)
        #self.genResultListWithDetail(self.manager.serverItems, resultList, 1)

        doc.save(out_word_file)

    #输出当前进度            
    def progress(self, subject, bFinish=False):
        self.progress_value += 1
        if self.progress_value >= self.progress_total or bFinish:
            self.progress_value = self.progress_total
        #输出进度
        #print("======= progress,%d%%,%s ============================"%(self.progress_value*100/self.progress_total, self.i18n(subject)))
        print("======= progress,%d%%,%s ============================"%(self.progress_value * 100 / self.progress_total, subject))

    #输出当前错误信息    
    def ERROR(self, msg):
        print("**** [ ERROR ] %s ************" % msg)

    #输出运行信息    
    def LOG_OUT(self, msg):
        print("---- [ LOG ] %s ------------ " %msg)

    def init(self, argv):
        if self.inited:
            return
        self.inited = True
        
        super(AssReport, self).init(argv)
        #攻击检查
        self.report = Report()
        self.report.list[0].type = self.i18n('攻击检查')

        self.progress_value = 0
        self.progress_total = 0

        for module in self.modules:
            module.init(argv)
        
    def run(self):
        print("====Check Start %s" % time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time())))
        self.setBaseInfo(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time())), 9)    #检测开始时间
        for module in self.modules:
            moduleName = str(module.__class__.__name__)
            try:
                module.run()
            except:
                #print ex.message
                traceback.print_exc()
                self.ERROR("Error Check " + module.__class__.__name__)

        self.updateResult()

        if self.result == -2:
            print "====Check End %s" % time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time()))
            self.progress("扫描完成", True)
            return

        print("==== Report Init =======================")
        if self.result == 1:    #低危
            self.setBaseInfo(u'低危', 7)
            self.setBaseInfo(u'您的应用检测结果为低危，希望您能根据建议及时修正风险项，如果有任何疑问请联系我们修正，建议您申请手工检测发现更多未知风险', 8)
        elif self.result == 2:  #中危
            self.setBaseInfo(u'中危', 7)
            self.setBaseInfo(u'您的应用检测结果为中危，请及时联系我们协助您修正风险项，建议您申请手工检测来确认其他未知风险', 8)
        elif self.result == 3:  #高危
            self.setBaseInfo(u'高危', 7)
            self.setBaseInfo(u'您的应用检测结果为高危，您的APP应用安全问题突出，请赶紧联系我们协助您修正风险项，建议您马上申请手工检测来确认其他未知风险，我们将为您量身定制一整套安全加固方案', 8)
        elif self.result == 0:  #安全
            self.setBaseInfo(u'安全', 7)
            self.setBaseInfo(u'恭喜您的应用检测结果为安全，自动化测试中没有发现可疑风险项问题，但建议您申请手工检测对该APP进行更加详尽的风险项检测，', 8)
        self.setBaseInfo(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time())), 10)   #检测结束时间

        self.saveReport(self.apk_file, moduleName)
        print "==== Report End ======================="

        print "====Check End %s" % time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time()))
        self.progress("扫描完成", True)

    def readReport(self, filename):
        reportStr = ass_base.read_file(filename)
        print "&&&&&&&&&&"
        d = AssDecoder().decode(reportStr)
        reportTmp = Report()
        for key, value in d.items():
            if key == 'basic':
                reportTmp.basic = Basic()
                for bkey, bvalue in value.items():
                    if bkey == 'appName':
                        reportTmp.basic.appName = bvalue
                    elif bkey == 'appVersion':
                        reportTmp.basic.appVersion = bvalue
                    elif bkey == 'packageName':
                        reportTmp.basic.packageName = bvalue
            elif key == 'result':
                reportTmp.result = Result()
                for rkey, rvalue in value.items():
                    if rkey == 'lastValue':
                        reportTmp.result.lastValue = rvalue
                    elif rkey == 'evaAdvice':
                        reportTmp.result.evaAdvice = rvalue
            elif key == 'list':
                reportTmp.list = [List()]
                for lkey, lvalue in value[0].items():
                    if lkey == 'type':
                        reportTmp.list[0].type = lvalue
                    elif lkey == 'items':
                        reportTmp.list[0].items = []
                        for i in range(0, len(lvalue)):
                            reportTmp.list[0].items.append(Items())
                            for ikey, ivalue in lvalue[i].items():
                                if ikey == 'checkMethod':
                                    reportTmp.list[0].items[i].checkMethod = ivalue
                                elif ikey == 'checkName':
                                    reportTmp.list[0].items[i].checkName = ivalue
                                elif ikey == 'checkResult':
                                    reportTmp.list[0].items[i].checkResult = ivalue
                                elif ikey == 'relatedCode':
                                    reportTmp.list[0].items[i].relatedCode = ivalue
                                elif ikey == 'resultComment':
                                    reportTmp.list[0].items[i].resultComment = ivalue

        for i in range(len(reportTmp.list)):
            for j in range(len(reportTmp.list[i].items)):
                print reportTmp.list[i].items[j].checkMethod + "*********"
                print reportTmp.list[i].items[j].checkName + "*********"
                print reportTmp.list[i].items[j].checkResult + "*********"
                print reportTmp.list[i].items[j].relatedCode + "*********"
                print reportTmp.list[i].items[j].resultComment + "*********"

if __name__=="__main__":
    AssReport().readReport('H:/AppSafeCheckLab/APKs/6b35d7e3-0571-44ba-acd5-1972209e0166.apk-report.json')
